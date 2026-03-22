#!/usr/bin/env python3
"""
Standalone Document Generation Script
Generates 25+ professional Word documents for custom electric engine order management system
Organized by domain: Business, Architecture, Systems Engineering, Design, Operations
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
import os

DOCS_DIR = Path(__file__).parent / 'docs'

def create_header(doc, title, doc_type, version='1.0', date=None):
    """Add document header with metadata table"""
    if date is None:
        date = datetime.now().strftime('%B %d, %Y')
    
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    cells[0].text = 'Document Type:'
    cells[1].text = doc_type
    
    cells = table.rows[1].cells
    cells[0].text = 'Version:'
    cells[1].text = version
    
    cells = table.rows[2].cells
    cells[0].text = 'Date:'
    cells[1].text = date
    
    # Make first column bold
    for row in table.rows:
        for cell in row.cells[:1]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    
    doc.add_paragraph()  # spacing


def add_section(doc, title, points):
    """Add a section with title and bullet points"""
    heading = doc.add_heading(title, level=2)
    
    for point in points:
        if isinstance(point, str):
            doc.add_paragraph(point, style='List Bullet')
        elif isinstance(point, tuple):
            # (text, is_heading) - for subsections
            if point[1]:
                doc.add_paragraph(point[0], style='List Number')
            else:
                doc.add_paragraph(point[0], style='List Bullet')


# Document definitions
DOCUMENTS = {
    'Business/BRD_Main.docx': {
        'title': 'Business Requirements Document - Custom Electric Engine Order System',
        'doc_type': 'BRD',
        'sections': [
            ('1. Executive Summary', [
                'This document outlines the business requirements for a comprehensive order management system designed to handle custom electric engine manufacturing orders from initial capture through final delivery and payment.'
            ]),
            ('2. Business Objectives', [
                'Streamline order capture and processing',
                'Reduce manual order handling by 80%',
                'Improve order accuracy and traceability',
                'Optimize supply chain management',
                'Enhance customer communication',
                'Accelerate order-to-delivery cycle',
                'Improve billing accuracy and payment collection'
            ]),
            ('3. Key Business Processes', [
                '3.1 Order Management - From initial customer request to order completion',
                '3.2 Supply Chain - Component sourcing and supplier management',
                '3.3 Manufacturing - Engine assembly and quality testing',
                '3.4 Financial Management - Billing, invoicing, and payment processing',
                '3.5 Customer Service - Communication and support throughout order lifecycle'
            ]),
            ('4. Success Metrics', [
                'Order processing time < 48 hours',
                '99.5% order accuracy rate',
                '95% on-time delivery',
                '98% payment collection within 30 days',
                'Customer satisfaction > 90%'
            ]),
            ('5. Scope and Constraints', [
                'Scope: End-to-end order management from intake through delivery',
                'Key Constraints: Must integrate with existing ERP systems, maintain data security and compliance'
            ])
        ]
    },

    'Business/Order_Management_Process.docx': {
        'title': 'Order Management Process Definition',
        'doc_type': 'Process',
        'sections': [
            ('1. Process Overview', [
                'The order management process handles all stages of customer orders from initial contact through post-delivery follow-up.'
            ]),
            ('2. Order Capture Phase', [
                '2.1 Customer Inquiry - Accept and log customer requests',
                '2.2 Initial Quote - Generate preliminary pricing and specifications',
                '2.3 Quote Review - Customer reviews and provides feedback',
                '2.4 Order Placement - Customer finalizes order with signed agreement and payment terms'
            ]),
            ('3. Order Processing Phase', [
                '3.1 Order Submission - Official order entry into system',
                '3.2 Order Decomposition - Break down complex orders into component specifications',
                '3.3 Spec Intake - Detailed technical specifications and custom requirements',
                '3.4 Spec Revision & Approval - Customer and engineering review cycles',
                '3.5 Bill of Materials (BOM) generation'
            ]),
            ('4. Order Fulfillment Phase', [
                '4.1 Supply planning and component sourcing',
                '4.2 Supplier selection and purchase orders',
                '4.3 Component receipt and quality inspection',
                '4.4 Assembly and manufacturing',
                '4.5 Quality testing and customer testing',
                '4.6 Delivery arrangement'
            ]),
            ('5. Order Closure Phase', [
                '5.1 Billing and invoicing',
                '5.2 Payment collection',
                '5.3 Product delivery',
                '5.4 Post-delivery support and warranty'
            ])
        ]
    },

    'Business/Supply_Chain_Process.docx': {
        'title': 'Supply Chain Management Process',
        'doc_type': 'Process',
        'sections': [
            ('1. Procurement Overview', [
                'This process defines how components are sourced, evaluated, and procured from suppliers.'
            ]),
            ('2. Supply Planning', [
                '2.1 Demand Forecasting - Predict component needs based on orders',
                '2.2 Inventory Management - Track stock levels and reorder points',
                '2.3 Supplier Database - Maintain current supplier information and capabilities'
            ]),
            ('3. Supplier Selection Process', [
                '3.1 RFQ (Request for Quote) Distribution',
                '3.2 Bidding and Evaluation',
                '3.3 Cost and Quality Analysis',
                '3.4 Supplier Selection and PO Creation',
                '3.5 Supplier Contract Management'
            ]),
            ('4. Component Receipt & Inspection', [
                '4.1 Incoming Shipment Tracking',
                '4.2 Quality Inspection and Testing',
                '4.3 Defect Management and Returns',
                '4.4 Inventory Receipt and Storage',
                '4.5 Component Allocation to Orders'
            ]),
            ('5. Supplier Payment Management', [
                '5.1 Invoice Matching and Validation',
                '5.2 Payment Processing',
                '5.3 Supplier Performance Tracking',
                '5.4 Dispute Resolution'
            ])
        ]
    },

    'Business/Billing_Payment_Process.docx': {
        'title': 'Billing and Payment Process',
        'doc_type': 'Process',
        'sections': [
            ('1. Process Overview', [
                'Defines the process for invoicing customers and collecting payments.'
            ]),
            ('2. Billing Strategy', [
                '2.1 50% deposit upon order confirmation',
                '2.2 30% upon BOM approval',
                '2.3 15% upon component receipt',
                '2.4 5% upon final delivery and customer acceptance'
            ]),
            ('3. Invoice Generation', [
                '3.1 Automatic invoice creation at billing milestones',
                '3.2 Invoice validation and approval',
                '3.3 Customer invoice delivery',
                '3.4 Payment terms enforcement (Net 30)'
            ]),
            ('4. Payment Collection', [
                '4.1 Payment receipt and allocation',
                '4.2 Payment reconciliation',
                '4.3 Late payment management',
                '4.4 Dispute handling and credit notes'
            ]),
            ('5. Financial Reporting', [
                '5.1 Revenue recognition',
                '5.2 Accounts receivable aging',
                '5.3 Collection rate reporting',
                '5.4 Financial forecasting'
            ])
        ]
    },

    'Business/Customer_Communication_Plan.docx': {
        'title': 'Customer Communication Plan',
        'doc_type': 'Communication',
        'sections': [
            ('1. Communication Objectives', [
                'Maintain transparency and regular updates throughout the order lifecycle'
            ]),
            ('2. Communication Touchpoints', [
                '2.1 Order Confirmation - Email with order details and timeline',
                '2.2 Spec Approval - Notification of approval status',
                '2.3 Manufacturing Start - Notification when assembly begins',
                '2.4 Quality Testing - Results and status updates',
                '2.5 Delivery Notification - Estimated delivery date and tracking',
                '2.6 Delivery Confirmation - Proof of delivery and follow-up support'
            ]),
            ('3. Escalation Procedures', [
                '3.1 Issue Reporting - Customer support channel',
                '3.2 Priority Assessment',
                '3.3 Resolution and Follow-up',
                '3.4 Executive escalation for critical issues'
            ]),
            ('4. Feedback and Satisfaction', [
                '4.1 Post-delivery satisfaction survey',
                '4.2 Quarterly business reviews',
                '4.3 Feedback incorporation into process improvements'
            ])
        ]
    },

    'Business/Project_Business_Case.docx': {
        'title': 'Engineering Firm Custom Electric Engine Order System - Business Case',
        'doc_type': 'BusinessCase',
        'sections': [
            ('1. Investment Summary', [
                'Estimated Investment: $500K - $750K',
                'Expected ROI: 150-200% over 3 years',
                'Payback Period: 18-24 months'
            ]),
            ('2. Current Pain Points', [
                'Manual order processing - 4-6 days per order',
                'Frequent errors in BOM and specifications',
                'Supplier coordination challenges',
                'Delayed billing and payment collection',
                'Poor visibility into order status'
            ]),
            ('3. Expected Benefits', [
                'Cost Savings: 30-40% reduction in processing labor',
                'Revenue Impact: Faster turnaround enables 20% more orders annually',
                'Margin Improvement: Better supplier negotiation saves 5-8%',
                'Working Capital Improvement: Faster payment collection improves cash flow by $2M+'
            ]),
            ('4. Risk Assessment', [
                '4.1 Implementation Risk: Medium - standard technology stack',
                '4.2 User Adoption: Medium - comprehensive training program planned',
                '4.3 Data Migration: Medium - legacy system data conversion required',
                '4.4 Mitigation Strategies: Phased roll-out, dedicated change management team'
            ])
        ]
    },

    'Business/End_to_End_Process_Flow.docx': {
        'title': 'End-to-End Order to Delivery Process Flow',
        'doc_type': 'Process Flow',
        'sections': [
            ('1. Process Overview', [
                'This document describes the complete order lifecycle from initial customer inquiry through final delivery, including parallel processing streams that can execute simultaneously.',
                'Process Duration Target: 20-30 business days (order capture to delivery)'
            ]),
            ('2. Phase 1: Order Capture & Intake (Days 1-2)', [
                '2.1 Customer submits inquiry with preliminary specifications',
                '2.2 Sales team creates quote and initial proposal',
                '2.3 Customer reviews and approves quote',
                '2.4 Order officially placed with signed agreement and initial payment (50% deposit)'
            ]),
            ('3. Phase 2: Order Processing & Parallel Setup (Days 2-3)', [
                '3.1 Order entered into system and assigned unique order ID',
                'PARALLEL STREAM A: Specification & Design',
                '  - Engineering team receives order specifications',
                '  - Detailed technical specifications developed',
                '  - Customer technical review meeting scheduled',
                'PARALLEL STREAM B: Supply Planning',
                '  - Supply chain team analyzes Bill of Materials requirements',
                '  - Supplier evaluation and selection process initiated',
                '  - RFQs distributed to qualified suppliers'
            ]),
            ('4. Phase 3: Design Review & Supplier Bidding (Days 3-5)', [
                'PARALLEL STREAM A: Specification Revision',
                '  - Customer receives draft specifications',
                '  - Technical review and revision cycles',
                '  - Final spec approval and signoff',
                '  - Generate Bill of Materials (BOM) - 30% invoice issued',
                'PARALLEL STREAM B: Supplier Selection',
                '  - Supplier bids received and evaluated',
                '  - Cost and quality analysis completed',
                '  - Preferred suppliers selected',
                '  - Purchase orders issued to suppliers'
            ]),
            ('5. Phase 4: Component Sourcing & Manufacturing Prep (Days 5-15)', [
                'PARALLEL STREAM A: Supply Chain Execution',
                '  - Suppliers manufacture/source components',
                '  - Delivery tracking and ETA management',
                '  - Components start arriving at warehouse',
                '  - Quality inspection and receiving process',
                '  - Defective items returned or replaced',
                '  - Components allocated to order batch',
                'PARALLEL STREAM B: Manufacturing Planning',
                '  - Manufacturing schedule created based on component arrival',
                '  - Assembly bay allocated and prepared',
                '  - Assembly team trained on custom specifications',
                '  - Testing protocols finalized',
                '  - Quality checkpoints defined'
            ]),
            ('6. Phase 5: Manufacturing & Testing (Days 12-22)', [
                'PARALLEL STREAM A: Assembly & Manufacturing',
                '  - Components picked from inventory',
                '  - Engine assembly begins (custom configuration)',
                '  - Assembly recorded in system with completion milestones',
                '  - In-process quality checks at each assembly stage',
                '  - Final assembly verification',
                'PARALLEL STREAM B: Documentation & Shipping Prep',
                '  - Assembly documentation prepared',
                '  - Test procedures and checklists finalized',
                '  - Generate final invoice - 15% payment collected upon component receipt',
                '  - Shipping label and documentation prepared',
                '  - Customer notified of manufacturing status'
            ]),
            ('7. Phase 6: Quality Testing & Validation (Days 20-25)', [
                'PARALLEL STREAM A: Engine Testing',
                '  - Performance testing executed',
                '  - Safety and compliance validation',
                '  - Test results documented',
                '  - Issues logged and resolved',
                '  - Final sign-off by QA team',
                'PARALLEL STREAM B: Customer Acceptance',
                '  - Customer notified of testing completion',
                '  - Optional customer witness testing',
                '  - Customer acceptance documentation received',
                '  - Final adjustments if required'
            ]),
            ('8. Phase 7: Delivery & Closure (Days 25-30)', [
                'PARALLEL STREAM A: Logistics',
                '  - Engine packaged for shipment',
                '  - Shipped via carrier with tracking',
                '  - Real-time delivery tracking provided to customer',
                '  - Delivery confirmation received',
                'PARALLEL STREAM B: Financial Closure',
                '  - Final invoice issued (5% final payment upon delivery)',
                '  - Payment collected',
                '  - Accounts receivable reconciliation',
                '  - Supplier invoices processed and paid'
            ]),
            ('9. Post-Delivery Activities (Days 25+)', [
                '9.1 Customer acknowledgment and sign-off on delivery',
                '9.2 Product warranty activation',
                '9.3 Customer satisfaction survey',
                '9.4 Post-delivery support and troubleshooting',
                '9.5 Performance metrics reported and archived',
                '9.6 Lessons learned captured for future improvements'
            ]),
            ('10. Critical Path Analysis', [
                'Critical Path = Longest sequence of dependent tasks',
                'Day 1-2: Order Capture',
                'Day 2-5: Design & Specification (can overlap with sourcing)',
                'Day 2-15: Component Sourcing (longest-lead supplier path)',
                'Day 12-22: Manufacturing (cannot start until components arrive)',
                'Day 20-25: Testing (must complete before shipment)',
                'Day 25-30: Delivery (fixed end point)',
                'Optimization: Start sourcing immediately while specs finalize to reduce overall cycle time'
            ]),
            ('11. Parallel Opportunities Summary', [
                'Specification & Design (Stream A) executes in parallel with Supply Planning (Stream B)',
                'Supplier Selection (Stream B) executes in parallel with Design Reviews (Stream A)',
                'Manufacturing Prep (Stream B) executes in parallel with Component Sourcing (Stream A)',
                'Quality Testing (Stream A) can overlap with Shipping Prep (Stream B)',
                'Logistics (Stream A) overlaps with Financial Closure (Stream B)',
                'Total parallelization saves approximately 5-7 days compared to sequential processing'
            ]),
            ('12. Key Milestones & Gating Items', [
                'Gate 1: Order Confirmation (after customer payment)',
                'Gate 2: Final Specifications Approved (before BOM finalization)',
                'Gate 3: All Components Received (before manufacturing start)',
                'Gate 4: Testing Passed (before shipment)',
                'Gate 5: Final Payment Received (before delivery)',
                'Gate 6: Delivery Confirmed (order complete)'
            ])
        ]
    },

    'Architecture/System_Architecture_Overview.docx': {
        'title': 'System Architecture Overview',
        'doc_type': 'Architecture',
        'sections': [
            ('1. Architecture Vision', [
                'A scalable, cloud-based, microservices-oriented architecture that enables rapid order processing and real-time visibility across the entire supply chain.'
            ]),
            ('2. Key Architecture Principles', [
                'Modularity - Loosely coupled, independently deployable services',
                'Scalability - Horizontal scaling for peak demand periods',
                'Resilience - Fault tolerance and graceful degradation',
                'Security - Defense in depth with role-based access control',
                'Interoperability - Standards-based APIs and data formats'
            ]),
            ('3. Core System Components', [
                '3.1 Order Management Service',
                '3.2 Supply Chain Management Service',
                '3.3 Inventory Management Service',
                '3.4 Manufacturing Planning Service',
                '3.5 Quality Management Service',
                '3.6 Billing & Payment Service',
                '3.7 Reporting & Analytics Service',
                '3.8 Customer Portal',
                '3.9 Integration Hub'
            ]),
            ('4. Technology Stack', [
                'Backend: Node.js/Java microservices',
                'Frontend: React/Vue.js',
                'Database: PostgreSQL (relational) + MongoDB (document)',
                'Message Queue: RabbitMQ/Kafka',
                'Cloud Platform: AWS/Azure',
                'Container Orchestration: Kubernetes',
                'API Gateway: Kong/AWS API Gateway'
            ]),
            ('5. Deployment Architecture', [
                '5.1 Development Environment',
                '5.2 Staging Environment',
                '5.3 Production Environment with multi-region failover',
                '5.4 Continuous Integration/Continuous Deployment (CI/CD) pipelines'
            ])
        ]
    },

    'Architecture/Data_Architecture.docx': {
        'title': 'Data Architecture',
        'doc_type': 'Architecture',
        'sections': [
            ('1. Data Strategy', [
                'Implement a hybrid data architecture combining relational and document databases for optimal performance and flexibility.'
            ]),
            ('2. Master Data Management', [
                '2.1 Customers - Customer accounts and profiles',
                '2.2 Suppliers - Supplier information and capabilities',
                '2.3 Components - Component catalog with specifications',
                '2.4 Products - Engine configurations and bill of materials templates',
                '2.5 Pricing - Dynamic pricing and cost models'
            ]),
            ('3. Transactional Data', [
                '3.1 Orders - Order headers, line items, and revision history',
                '3.2 Specifications - Technical specs and custom requirements',
                '3.3 Supply Orders - Purchase orders and supplier orders',
                '3.4 Inventory - Stock levels, movements, and allocations',
                '3.5 Manufacturing - Assembly logs and testing results',
                '3.6 Billing & Payments - Invoices, payments, and reconciliation'
            ]),
            ('4. Analytics Data', [
                '4.1 Data Warehouse - Aggregated data for reporting and analytics',
                '4.2 Data Lake - Raw data retention for advanced analytics',
                '4.3 Real-time Dashboards - Event streaming and aggregation',
                '4.4 Metrics and KPIs - Business performance indicators'
            ]),
            ('5. Data Governance', [
                '5.1 Data Quality Standards',
                '5.2 Data Classification',
                '5.3 Retention Policies',
                '5.4 Access Control and Audit Trails'
            ])
        ]
    },

    'Architecture/Integration_Architecture.docx': {
        'title': 'Integration Architecture',
        'doc_type': 'Architecture',
        'sections': [
            ('1. Integration Strategy', [
                'Implement an event-driven integration architecture using APIs and message queues for real-time data synchronization.'
            ]),
            ('2. External System Integrations', [
                '2.1 ERP System - Financial data and accounting integration',
                '2.2 Email/Communication - Customer and internal notifications',
                '2.3 Payment Gateway - Credit card and ACH processing',
                '2.4 Shipping Carriers - Real-time tracking and label generation',
                '2.5 Supplier Systems - EDI and API integration'
            ]),
            ('3. API Gateway', [
                '3.1 RESTful APIs for all microservices',
                '3.2 API versioning and backward compatibility',
                '3.3 Rate limiting and throttling',
                '3.4 Request/response logging and monitoring',
                '3.5 API documentation (OpenAPI/Swagger)'
            ]),
            ('4. Message Queue Architecture', [
                '4.1 Event Publishing - Services publish domain events',
                '4.2 Event Subscription - Services subscribe to relevant events',
                '4.3 Dead Letter Handling - Failed messages and retry logic',
                '4.4 Message Ordering and Idempotency'
            ]),
            ('5. Integration Patterns', [
                '5.1 Request-Reply: Synchronous API calls',
                '5.2 Publish-Subscribe: Event-driven async communication',
                '5.3 Saga Pattern: Distributed transactions',
                '5.4 Circuit Breaker: Fault tolerance'
            ])
        ]
    },

    'Architecture/Security_Architecture.docx': {
        'title': 'Security Architecture',
        'doc_type': 'Architecture',
        'sections': [
            ('1. Security Principles', [
                'Defense in depth, least privilege, zero trust architecture'
            ]),
            ('2. Authentication & Authorization', [
                '2.1 Multi-factor authentication (MFA) for all users',
                '2.2 OAuth 2.0/OIDC for API authentication',
                '2.3 Role-Based Access Control (RBAC)',
                '2.4 Session management and timeout policies',
                '2.5 API key management and rotation'
            ]),
            ('3. Data Protection', [
                '3.1 Encryption in transit (TLS 1.3)',
                '3.2 Encryption at rest (AES-256)',
                '3.3 Database encryption and transparent data encryption',
                '3.4 Secure key management (KMS)',
                '3.5 PCI DSS compliance for payment data'
            ]),
            ('4. Network Security', [
                '4.1 VPC isolation and network segmentation',
                '4.2 WAF (Web Application Firewall)',
                '4.3 DDoS protection',
                '4.4 Intrusion detection/prevention',
                '4.5 VPN for remote access'
            ]),
            ('5. Compliance & Auditing', [
                '5.1 GDPR and data privacy compliance',
                '5.2 Audit logging and traceability',
                '5.3 Regular security assessments and penetration testing',
                '5.4 Incident response plan',
                '5.5 Security awareness training'
            ])
        ]
    },

    'Architecture/Scalability_Performance.docx': {
        'title': 'Scalability and Performance Architecture',
        'doc_type': 'Architecture',
        'sections': [
            ('1. Scalability Strategy', [
                'Implement horizontal scaling with stateless services and distributed caching.'
            ]),
            ('2. Load Balancing', [
                '2.1 Multi-tier load balancing (DNS, ALB, internal)',
                '2.2 Session affinity and sticky sessions',
                '2.3 Connection pooling and draining',
                '2.4 Auto-scaling policies based on CPU and memory'
            ]),
            ('3. Caching Strategy', [
                '3.1 Redis for distributed caching',
                '3.2 Cache invalidation and TTL policies',
                '3.3 CDN for static content delivery',
                '3.4 Client-side caching',
                '3.5 Cache warming for critical data'
            ]),
            ('4. Database Performance', [
                '4.1 Query optimization and indexing strategies',
                '4.2 Read replicas for scaling read operations',
                '4.3 Connection pooling',
                '4.4 Sharding for large tables',
                '4.5 Regular maintenance and statistics'
            ]),
            ('5. Performance Targets', [
                'API response time: < 200ms at p95',
                'Page load time: < 2 seconds',
                'System throughput: 1000+ TPS',
                'Concurrent users: 10,000+'
            ])
        ]
    },

    'Architecture/Disaster_Recovery.docx': {
        'title': 'Disaster Recovery and Business Continuity',
        'doc_type': 'Architecture',
        'sections': [
            ('1. RTO and RPO Targets', [
                'Recovery Time Objective (RTO): 1 hour',
                'Recovery Point Objective (RPO): 15 minutes'
            ]),
            ('2. Backup Strategy', [
                '2.1 Continuous replication to secondary region',
                '2.2 Daily incremental backups with weekly full backups',
                '2.3 Test restoration procedures quarterly',
                '2.4 Off-site backup storage with encryption',
                '2.5 Database point-in-time recovery capability'
            ]),
            ('3. Redundancy Architecture', [
                '3.1 Multi-region deployment',
                '3.2 Database replication with failover',
                '3.3 Redundant network connections',
                '3.4 DNS failover',
                '3.5 Load balancer redundancy'
            ]),
            ('4. Disaster Recovery Procedures', [
                '4.1 Incident detection and alerting',
                '4.2 Failover procedures',
                '4.3 Data consistency verification',
                '4.4 Communication plan',
                '4.5 Failback procedures'
            ]),
            ('5. Continuity Operations', [
                '5.1 Manual order processing procedures',
                '5.2 Customer communication templates',
                '5.3 Temporary data entry systems',
                '5.4 Supply chain contingency coordination'
            ])
        ]
    },

    'SystemsEngineering/Systems_Engineering_Plan.docx': {
        'title': 'Systems Engineering Plan',
        'doc_type': 'SystemsEngineering',
        'sections': [
            ('1. SE Organization', [
                'Systems Engineering Lead: Reports to Program Manager',
                'Requirements Engineering Team',
                'Integration & Test Team',
                'Configuration Management Team'
            ]),
            ('2. SE Process Framework', [
                '2.1 Requirements Development - Define and analyze requirements',
                '2.2 Systems Design - Develop and refine design solutions',
                '2.3 Implementation - Develop and build system components',
                '2.4 Integration - Assemble and integrate components',
                '2.5 Verification & Validation - Ensure requirements are met',
                '2.6 Operations & Support - Deploy and maintain system'
            ]),
            ('3. Key Processes', [
                '3.1 Requirements Management Process',
                '3.2 Configuration Management Process',
                '3.3 Design Review Process',
                '3.4 Integration & Test Process',
                '3.5 Risk Management Process',
                '3.6 Change Management Process'
            ]),
            ('4. Governance & Reviews', [
                '4.1 Weekly SE status meetings',
                '4.2 Monthly design reviews',
                '4.3 Quarterly milestone reviews',
                '4.4 Executive steering committee reviews'
            ]),
            ('5. Success Criteria', [
                'All system requirements traced to design',
                '100% requirement coverage in test cases',
                'Defect closure rate > 95%',
                'Schedule and budget performance'
            ])
        ]
    },

    'SystemsEngineering/Requirements_Management_Plan.docx': {
        'title': 'Requirements Management Plan',
        'doc_type': 'SystemsEngineering',
        'sections': [
            ('1. Purpose', [
                'Establish processes for collecting, analyzing, and managing system requirements throughout the project lifecycle.'
            ]),
            ('2. Requirements Types', [
                '2.1 Functional Requirements - What the system must do',
                '2.2 Non-Functional Requirements - Performance, security, scalability',
                '2.3 Operational Requirements - Deployment, support, maintenance',
                '2.4 Compliance Requirements - Legal, regulatory, standards'
            ]),
            ('3. Requirement Lifecycle', [
                '3.1 Elicitation - Gather from stakeholders',
                '3.2 Analysis - Assess feasibility and completeness',
                '3.3 Specification - Document with acceptance criteria',
                '3.4 Verification - Ensure requirements are understood',
                '3.5 Traceability - Link requirements to design and tests',
                '3.6 Change Management - Control requirement changes'
            ]),
            ('4. Requirements Document Structure', [
                'Requirement ID, Title, Description, Priority, Acceptance Criteria, Dependencies'
            ]),
            ('5. Requirements Traceability Matrix', [
                'Maintain traceability from requirements to design specifications, code, and test cases'
            ])
        ]
    },

    'SystemsEngineering/Configuration_Management_Plan.docx': {
        'title': 'Configuration Management Plan',
        'doc_type': 'SystemsEngineering',
        'sections': [
            ('1. Purpose', [
                'Establish configuration items, baselines, and change control procedures.'
            ]),
            ('2. Configuration Items (CIs)', [
                '2.1 Source Code - All application code and libraries',
                '2.2 Documentation - Requirements, design, and user guides',
                '2.3 Data - Database schemas and seed data',
                '2.4 Infrastructure - Configuration files and scripts',
                '2.5 Build Artifacts - Compiled binaries and packages'
            ]),
            ('3. Configuration Baselines', [
                '3.1 Requirements Baseline - Approved system requirements',
                '3.2 Design Baseline - Approved design specification',
                '3.3 Build Baseline - Tested and approved builds',
                '3.4 Release Baseline - Production-ready releases'
            ]),
            ('4. Change Control Process', [
                '4.1 Change Request Submission',
                '4.2 Impact Assessment',
                '4.3 Change Review Board (CRB) approval',
                '4.4 Implementation and verification',
                '4.5 Baseline update and communication'
            ]),
            ('5. Version Control', [
                '5.1 Git branching strategy (Git Flow)',
                '5.2 Tag releases with semantic versioning',
                '5.3 Maintain change logs',
                '5.4 Release notes generation'
            ])
        ]
    },

    'SystemsEngineering/Verification_Validation_Plan.docx': {
        'title': 'Verification & Validation Plan',
        'doc_type': 'SystemsEngineering',
        'sections': [
            ('1. V&V Strategy', [
                'Verification: Does the system implement the requirements correctly?',
                'Validation: Does the system meet business needs and user expectations?'
            ]),
            ('2. Test Levels', [
                '2.1 Unit Testing - Individual component testing',
                '2.2 Integration Testing - Component interaction testing',
                '2.3 System Testing - End-to-end functionality testing',
                '2.4 User Acceptance Testing - Customer validation',
                '2.5 Regression Testing - Verify no defects introduced'
            ]),
            ('3. Test Types', [
                '3.1 Functional Testing - Requirements validation',
                '3.2 Performance Testing - Load and stress testing',
                '3.3 Security Testing - Vulnerability assessment',
                '3.4 Usability Testing - User experience validation',
                '3.5 Compatibility Testing - Platform compatibility'
            ]),
            ('4. Test Coverage Targets', [
                'Unit test coverage: > 85%',
                'Requirement coverage: 100%',
                'Critical path coverage: 100%'
            ]),
            ('5. Defect Management', [
                '5.1 Defect logging and severity classification',
                '5.2 Root cause analysis',
                '5.3 Fix verification and regression testing',
                '5.4 Metrics and trend analysis'
            ])
        ]
    },

    'Design/Database_Design.docx': {
        'title': 'Database Design Document',
        'doc_type': 'Design',
        'sections': [
            ('1. Database Architecture', [
                'PostgreSQL for transactional data and relational integrity',
                'MongoDB for document storage (specifications, attachments)',
                'Redis for real-time caching and message queue',
                'Elasticsearch for full-text search and analytics'
            ]),
            ('2. Core Tables', [
                '2.1 customers - Customer accounts',
                '2.2 orders - Order headers and metadata',
                '2.3 order_items - Order line items',
                '2.4 order_specifications - Custom technical specifications',
                '2.5 suppliers - Supplier information',
                '2.6 components - Component catalog',
                '2.7 supply_orders - Purchase orders',
                '2.8 inventory - Stock tracking',
                '2.9 manufacturing_batches - Assembly batches',
                '2.10 quality_tests - Test results',
                '2.11 invoices - Billing documents',
                '2.12 payments - Payment records'
            ]),
            ('3. Key Design Decisions', [
                '3.1 Soft deletes for audit trail preservation',
                '3.2 Audit timestamps (created_at, updated_at, deleted_at)',
                '3.3 Status enumerations for workflow states',
                '3.4 JSON fields for flexible attributes',
                '3.5 Partitioning strategy for large tables by date'
            ]),
            ('4. Indexing Strategy', [
                '4.1 Primary keys on all tables',
                '4.2 Foreign key indexes for referential integrity',
                '4.3 Composite indexes on frequently filtered columns',
                '4.4 Partial indexes for status queries'
            ]),
            ('5. Data Migration', [
                '5.1 Legacy system data mapping',
                '5.2 Data validation and cleansing',
                '5.3 Parallel run procedures',
                '5.4 Cutover plan and rollback procedures'
            ])
        ]
    },

    'Design/API_Design.docx': {
        'title': 'API Design Specification',
        'doc_type': 'Design',
        'sections': [
            ('1. API Principles', [
                'RESTful design with HTTP verbs',
                'JSON request/response format',
                'Versioning through URL path (v1, v2)',
                'Pagination for list endpoints',
                'Comprehensive error responses'
            ]),
            ('2. Core API Resources', [
                '2.1 /api/v1/orders - Order management',
                '2.2 /api/v1/customers - Customer management',
                '2.3 /api/v1/specifications - Technical specifications',
                '2.4 /api/v1/suppliers - Supplier management',
                '2.5 /api/v1/components - Component catalog',
                '2.6 /api/v1/supply-orders - Purchase orders',
                '2.7 /api/v1/inventory - Inventory management',
                '2.8 /api/v1/manufacturing - Manufacturing tracking',
                '2.9 /api/v1/invoices - Billing',
                '2.10 /api/v1/reports - Reporting and analytics'
            ]),
            ('3. Authentication & Authorization', [
                '3.1 OAuth 2.0 with JWT tokens',
                '3.2 Scope-based permissions',
                '3.3 Rate limiting (1000 requests/hour)',
                '3.4 API key management for integrations'
            ]),
            ('4. Request/Response Examples', [
                '4.1 Sample request bodies with required fields',
                '4.2 Response schemas and status codes',
                '4.3 Error response format with error codes',
                '4.4 Pagination parameters and metadata'
            ]),
            ('5. Documentation & Tools', [
                'OpenAPI 3.0 specification',
                'Swagger UI for interactive documentation',
                'Postman collection for testing'
            ])
        ]
    },

    'Design/UI_Order_Management.docx': {
        'title': 'Order Management UI Design',
        'doc_type': 'Design',
        'sections': [
            ('1. UI Philosophy', [
                'Simple, intuitive workflows with minimal navigation',
                'Real-time status visibility',
                'Quick actions and bulk operations',
                'Responsive design for mobile and tablet'
            ]),
            ('2. Key Pages/Screens', [
                '2.1 Order List - Searchable, filterable order list with status colors',
                '2.2 Order Detail - Complete order information with timeline',
                '2.3 Order Creation Wizard - Step-by-step order entry',
                '2.4 Specification Management - Spec creation and revision workflow',
                '2.5 BOM View - Bill of materials with component details',
                '2.6 Order Tracking - Customer-facing order status and tracking',
                '2.7 Supply Planning - Component sourcing and supplier management'
            ]),
            ('3. Design Elements', [
                '3.1 Color scheme - Blue primary, red for alerts, green for success',
                '3.2 Typography - Clear hierarchy with readable fonts',
                '3.3 Icons - Consistent icon set for actions and status',
                '3.4 Forms - Inline validation with helpful error messages',
                '3.5 Modals - Confirmation dialogs for critical actions'
            ]),
            ('4. Accessibility', [
                'WCAG 2.1 AA compliance',
                'Keyboard navigation support',
                'Screen reader compatible',
                'High contrast mode support'
            ]),
            ('5. Performance', [
                'Page load time < 2 seconds',
                'Smooth animations (60 FPS)',
                'Lazy loading for large lists',
                'Progressive enhancement for connectivity'
            ])
        ]
    },

    'Design/Dashboard_Design.docx': {
        'title': 'Executive Dashboard Design',
        'doc_type': 'Design',
        'sections': [
            ('1. Dashboard Purpose', [
                'Provide real-time KPI visibility to executive leadership'
            ]),
            ('2. Key Metrics', [
                '2.1 Orders - Total, pending, in-progress, completed',
                '2.2 Revenue - YTD, monthly, by customer',
                '2.3 Performance - On-time delivery %, cycle time',
                '2.4 Supply Chain - Supplier on-time %, defect rate',
                '2.5 Financial - Receivables aging, payment rate',
                '2.6 Quality - Defect rate, rework rate, testing pass rate'
            ]),
            ('3. Visualizations', [
                '3.1 KPI cards for key metrics',
                '3.2 Trend charts for historical performance',
                '3.3 Pie charts for distribution (e.g., orders by status)',
                '3.4 Heatmaps for supplier performance',
                '3.5 Gauges for targets and actual performance'
            ]),
            ('4. Interactivity', [
                '4.1 Drill-down capability to underlying data',
                '4.2 Date range selection',
                '4.3 Filtering by customer, supplier, region',
                '4.4 Exportable reports in PDF and Excel'
            ]),
            ('5. Refresh Rates', [
                'Real-time updates for critical metrics',
                'Hourly refresh for transaction metrics',
                'Daily refresh for historical analysis'
            ])
        ]
    },

    'Design/Reporting_Design.docx': {
        'title': 'Reporting System Design',
        'doc_type': 'Design',
        'sections': [
            ('1. Report Types', [
                '2.1 Operational Reports - Daily order and supply status',
                '2.2 Financial Reports - Revenue, receivables, expense analysis',
                '2.3 Quality Reports - Defect, rework, and testing metrics',
                '2.4 Supplier Performance - On-time, quality, cost analysis',
                '2.5 Customer Reports - Order history, spending, satisfaction'
            ]),
            ('2. Report Delivery', [
                '2.1 Scheduled reports (daily, weekly, monthly)',
                '2.2 Ad-hoc report generation',
                '2.3 Email distribution',
                '2.4 Portal access and archival'
            ]),
            ('3. Report Format Options', [
                'PDF for printing and archival',
                'Excel for further analysis',
                'HTML for web viewing',
                'JSON data export for integration'
            ]),
            ('4. Business Intelligence Integration', [
                '4.1 Data warehouse for aggregation',
                '4.2 BI tool integration (Tableau, Power BI)',
                '4.3 Custom metric definitions',
                '4.4 Self-service report builder'
            ]),
            ('5. Data Sources', [
                '5.1 Transactional database (real-time)',
                '5.2 Data warehouse (aggregated, historized)',
                '5.3 Event streams (real-time metrics)',
                '5.4 External data (market prices, supplier data)'
            ])
        ]
    },

    'Design/Mobile_App_Design.docx': {
        'title': 'Mobile Application Design',
        'doc_type': 'Design',
        'sections': [
            ('1. Mobile App Strategy', [
                'Native iOS and Android apps for optimal UX',
                'Focus on order tracking and status notifications',
                'Simple, touch-optimized interfaces',
                'Works offline for critical functionality'
            ]),
            ('2. Core Features', [
                '2.1 Order Search - Find orders by ID or date',
                '2.2 Order Tracking - Real-time status and timeline',
                '2.3 Notifications - Push alerts for order milestones',
                '2.4 Documents - View quotes, invoices, shipping labels',
                '2.5 Messaging - In-app chat with support',
                '2.6 Signature Capture - For delivery confirmation'
            ]),
            ('3. Design Patterns', [
                '3.1 Tab navigation for main sections',
                '3.2 Swipe gestures for quick navigation',
                '3.3 FAB (Floating Action Button) for primary action',
                '3.4 Bottom sheets for quick actions',
                '3.5 Native OS conventions (iOS Human Interface, Material Design)'
            ]),
            ('4. Performance & Storage', [
                '4.1 Minimal app size < 50MB',
                '4.2 Local data caching',
                '4.3 Offline-first architecture',
                '4.4 Battery-efficient background sync'
            ]),
            ('5. Security', [
                '5.1 Local data encryption',
                '5.2 Certificate pinning',
                '5.3 Biometric authentication',
                '5.4 Session timeout and logout on app close'
            ])
        ]
    },

    'Design/Integration_Patterns.docx': {
        'title': 'Integration Design Patterns',
        'doc_type': 'Design',
        'sections': [
            ('1. Synchronous Integration Patterns', [
                '1.1 Request-Reply - Direct API calls with response',
                '1.2 Query/Response - CQS pattern with read replicas',
                '1.3 Remote Procedure Call - RPC-style method invocation'
            ]),
            ('2. Asynchronous Integration Patterns', [
                '2.1 Publish-Subscribe - Event broadcasting',
                '2.2 Event Sourcing - Event-based state management',
                '2.3 Message Queue - Async message processing',
                '2.4 Polling - Consumer pulls data periodically'
            ]),
            ('3. Distribution Patterns', [
                '3.1 Service Mesh - Microservice communication layer',
                '3.2 API Gateway - Single entry point for clients',
                '3.3 Backend for Frontend (BFF) - Optimized service layers',
                '3.4 Event Bus - Central event distribution'
            ]),
            ('4. Reliability Patterns', [
                '4.1 Retry Logic - Exponential backoff',
                '4.2 Circuit Breaker - Prevent cascading failures',
                '4.3 Bulkhead - Isolate resources',
                '4.4 Timeout - Prevent hanging operations'
            ]),
            ('5. Example Integration Flows', [
                'Order submission triggers supply planning event',
                'Component receipt triggers inventory update and manufacturing notification',
                'Quality test completion triggers billing event',
                'Payment received updates customer credit status'
            ])
        ]
    },

    'Design/Error_Handling_Logging.docx': {
        'title': 'Error Handling & Logging Design',
        'doc_type': 'Design',
        'sections': [
            ('1. Error Handling Strategy', [
                'Graceful degradation - System continues with reduced functionality',
                'Fail-safe defaults - Safe defaults on critical failures',
                'User-friendly error messages',
                'Automatic recovery where possible'
            ]),
            ('2. Error Classification', [
                '2.1 Critical - System down, data loss risk',
                '2.2 Major - Feature unavailable, workaround available',
                '2.3 Minor - Cosmetic issues, no functional impact',
                '2.4 Info - Successfully completed operations'
            ]),
            ('3. Logging Strategy', [
                '3.1 Structured logging with JSON format',
                '3.2 Log levels - DEBUG, INFO, WARN, ERROR, CRITICAL',
                '3.3 Correlation IDs for request tracing',
                '3.4 Performance metrics and timing information',
                '3.5 Sensitive data masking'
            ]),
            ('4. Log Aggregation', [
                '4.1 ELK Stack (Elasticsearch, Logstash, Kibana)',
                '4.2 Centralized log storage',
                '4.3 Real-time search and analysis',
                '4.4 Alert thresholds for error rates',
                '4.5 Log retention policy (7 years for compliance)'
            ]),
            ('5. Monitoring & Alerting', [
                '5.1 Application Performance Monitoring (APM)',
                '5.2 Error rate thresholds',
                '5.3 SLA violation alerts',
                '5.4 Automated incident escalation'
            ])
        ]
    },

    'Operations/Operations_Manual.docx': {
        'title': 'System Operations Manual',
        'doc_type': 'Operations',
        'sections': [
            ('1. System Overview', [
                'The custom electric engine order management system is a cloud-based SaaS application deployed on AWS.'
            ]),
            ('2. Daily Operations', [
                '2.1 Morning health checks - Verify all services running',
                '2.2 Monitor dashboards - Check error rates and performance',
                '2.3 Process overnight batch jobs - Data aggregation and reporting',
                '2.4 Review alerts - Investigate and resolve issues',
                '2.5 End-of-day reconciliation - Financial and order data'
            ]),
            ('3. Access Management', [
                '3.1 User provisioning process - Approval and onboarding',
                '3.2 Password policies - Complexity, expiration, MFA',
                '3.3 Access revocation - Offboarding and role changes',
                '3.4 Audit logging - Track all access and changes'
            ]),
            ('4. System Health Monitoring', [
                '4.1 CPU and memory utilization < 80%',
                '4.2 Disk space utilization < 85%',
                '4.3 Database replication lag < 5 seconds',
                '4.4 API response time < 200ms',
                '4.5 Error rate < 0.1%'
            ]),
            ('5. Escalation Procedures', [
                '5.1 Level 1 - Ops team - Routine monitoring and alerts',
                '5.2 Level 2 - Engineering - Investigation and fixes',
                '5.3 Level 3 - Architecture - Complex system issues',
                '5.4 Level 4 - CTO - Critical system failures'
            ])
        ]
    },

    'Operations/Deployment_Guide.docx': {
        'title': 'Deployment Guide',
        'doc_type': 'Operations',
        'sections': [
            ('1. Deployment Pipeline', [
                '1.1 Development - Local development and testing',
                '1.2 Staging - Pre-production environment identical to production',
                '1.3 Production - Live customer environment',
                '1.4 Canary - Gradual rollout to percentage of traffic'
            ]),
            ('2. Release Process', [
                '2.1 Tag release in Git',
                '2.2 Automated tests must pass',
                '2.3 Build Docker images',
                '2.4 Push to image registry',
                '2.5 Deploy to staging for validation',
                '2.6 Smoke tests on staging',
                '2.7 Production deployment approval'
            ]),
            ('3. Production Deployment', [
                '3.1 Deploy during maintenance window (2-4 AM)',
                '3.2 Blue-green deployment for zero downtime',
                '3.3 Database migrations pre-tested',
                '3.4 Health checks after deployment',
                '3.5 Monitor for errors and performance degradation'
            ]),
            ('4. Rollback Procedures', [
                '4.1 Quick rollback if critical errors detected',
                '4.2 Database rollback to pre-migration state',
                '4.3 Customer communication of issue and fix'
            ]),
            ('5. Release Notes', [
                '5.1 Features added and improved',
                '5.2 Bugs fixed',
                '5.3 Known issues and workarounds',
                '5.4 Upgrade instructions if needed'
            ])
        ]
    },

    'Operations/Monitoring_Strategy.docx': {
        'title': 'Monitoring & Alerting Strategy',
        'doc_type': 'Operations',
        'sections': [
            ('1. Monitoring Tools', [
                'Application Monitoring: Datadog/New Relic',
                'Infrastructure Monitoring: CloudWatch/Azure Monitor',
                'Database Monitoring: RDS Performance Insights',
                'Log Monitoring: ELK Stack/Splunk',
                'Synthetic Monitoring: Uptime Robot/Pingdom'
            ]),
            ('2. Key Metrics to Monitor', [
                '2.1 Availability - Uptime % (target: 99.95%)',
                '2.2 Performance - Response time p95, p99',
                '2.3 Throughput - Requests per second, transactions',
                '2.4 Error Rate - Errors per minute, error types',
                '2.5 Resource Utilization - CPU, memory, disk, network',
                '2.6 Business Metrics - Orders processed, revenue, customers'
            ]),
            ('3. Alert Thresholds', [
                '3.1 Critical - Page ops immediately (availability < 99%, error rate > 1%)',
                '3.2 Warning - Alert ops team (availability < 99.5%, error rate > 0.5%)',
                '3.3 Info - Log for analysis (high response time, resource warnings)'
            ]),
            ('4. Alerting Channels', [
                '4.1 SMS and phone for critical alerts',
                '4.2 Email for warnings',
                '4.3 Slack for information alerts',
                '4.4 Automated tickets for investigation'
            ]),
            ('5. Alert Fatigue Mitigation', [
                '5.1 Intelligent alert aggregation',
                '5.2 Correlation of related alerts',
                '5.3 Maintenance windows and suppressions',
                '5.4 Regular alert threshold tuning'
            ])
        ]
    },

    'Operations/Incident_Management.docx': {
        'title': 'Incident Management Procedures',
        'doc_type': 'Operations',
        'sections': [
            ('1. Incident Definition', [
                'Unplanned interruption or significant reduction in service quality affecting users'
            ]),
            ('2. Incident Severity Levels', [
                'P1 - Critical: System down, all users affected, >$10k/hr impact',
                'P2 - High: Major feature unavailable, >50% users affected',
                'P3 - Medium: Workaround available, <50% users affected',
                'P4 - Low: Cosmetic issues, no functional impact'
            ]),
            ('3. Response Procedures', [
                '3.1 Alert received - Page appropriate on-call engineer',
                '3.2 Initial assessment - Confirm incident, estimate impact',
                '3.3 Incident commander assigned - Coordinates response',
                '3.4 Customer notification - Communicate status and ETA',
                '3.5 Investigation and diagnosis',
                '3.6 Implementation of fix',
                '3.7 Validation and recovery',
                '3.8 Post-incident review and action items'
            ]),
            ('4. Communication Protocols', [
                '4.1 War room - Slack channel or conference bridge',
                '4.2 Status updates - Every 15 minutes during incident',
                '4.3 Customer updates - Via status page and email',
                '4.4 Post-incident - Blameless post-mortem within 48 hours'
            ]),
            ('5. Escalation Matrix', [
                'Tier 1 Escalation: Engineering lead at 15 minutes',
                'Tier 2 Escalation: VP Engineering at 30 minutes',
                'Tier 3 Escalation: CTO/VP Operations at 1 hour'
            ])
        ]
    },

    'Operations/User_Training_Manual.docx': {
        'title': 'User Training Manual',
        'doc_type': 'Operations',
        'sections': [
            ('1. Training Objectives', [
                'Enable users to effectively use the system for their daily work',
                'Minimize support requests and system misuse',
                'Ensure consistent process compliance'
            ]),
            ('2. Role-Based Training', [
                '2.1 Sales User - Order creation, quoting, customer interaction',
                '2.2 Operations User - Order processing, supply planning, manufacturing',
                '2.3 Finance User - Billing, payment processing, reporting',
                '2.4 Admin User - User management, system configuration',
                '2.5 Executive User - Dashboard, reporting, KPI monitoring'
            ]),
            ('3. Training Delivery', [
                '3.1 Live instructor-led sessions for initial training',
                '3.2 Video recordings for on-demand learning',
                '3.3 Quick reference guides for common tasks',
                '3.4 In-system help and tooltips',
                '3.5 Hands-on sandbox environment for practice'
            ]),
            ('4. Training Schedule', [
                '4.1 Pre-launch training - 2-3 weeks before go-live',
                '4.2 Refresher training - Quarterly updates',
                '4.3 Ad-hoc training - New hires and feature updates'
            ]),
            ('5. Success Criteria', [
                '> 90% user attendance at training',
                '> 85% pass rate on assessment tests',
                '< 5 support tickets per user in first month'
            ])
        ]
    },

    'Operations/System_Maintenance.docx': {
        'title': 'System Maintenance Schedule',
        'doc_type': 'Operations',
        'sections': [
            ('1. Maintenance Windows', [
                'Planned maintenance: 2-4 AM Sundays UTC',
                'Emergency maintenance: As needed with 1 hour notice to customers',
                'No maintenance during fiscal month-end (last 3 days)'
            ]),
            ('2. Regular Maintenance Tasks', [
                '2.1 Database Optimization - Weekly query optimization and index maintenance',
                '2.2 Security Patching - Monthly OS and library updates',
                '2.3 Certificate Renewal - 60 days before expiration',
                '2.4 Log Rotation - Archive logs quarterly',
                '2.5 Backup Verification - Monthly restore tests',
                '2.6 Performance Tuning - Monthly review and optimization'
            ]),
            ('3. Seasonal Maintenance', [
                '3.1 Q1 - Full security audit',
                '3.2 Q2 - Database migration staging',
                '3.3 Q3 - Infrastructure capacity planning',
                '3.4 Q4 - Year-end data archive and optimization'
            ]),
            ('4. Change Management', [
                '4.1 All changes require post-maintenance approval',
                '4.2 High-risk changes tested in staging',
                '4.3 Rollback procedures documented',
                '4.4 Communication to stakeholders'
            ]),
            ('5. Documentation', [
                '5.1 Maintenance logs for audit trail',
                '5.2 Performance reports before/after',
                '5.3 Issues encountered and resolutions',
                '5.4 Action items for next cycle'
            ])
        ]
    }
}


def generate_all_docs():
    """Generate all Word documents"""
    # Create directory structure
    for domain in ['Business', 'Architecture', 'SystemsEngineering', 'Design', 'Operations']:
        domain_dir = DOCS_DIR / domain
        domain_dir.mkdir(parents=True, exist_ok=True)
    
    count = 0
    for filename, config in DOCUMENTS.items():
        try:
            # Create document
            doc = Document()
            
            # Add header
            create_header(
                doc, 
                config['title'], 
                config['doc_type'],
                version='1.0',
                date=datetime.now().strftime('%B %d, %Y')
            )
            
            # Add sections
            for section_title, points in config['sections']:
                add_section(doc, section_title, points)
            
            # Save document
            filepath = DOCS_DIR / filename
            filepath.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(filepath))
            
            print(f'✓ Generated: {filename}')
            count += 1
            
        except Exception as e:
            print(f'✗ Failed to generate {filename}: {str(e)}')
    
    print(f'\n✓ Successfully generated {count} documents in {DOCS_DIR}')


if __name__ == '__main__':
    generate_all_docs()
